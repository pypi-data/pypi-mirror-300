from docx import Document
from docx.shared import Inches
from typing import List
import folium
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from field_manager_api.methods.get_methods import Method


class WordDocReport:
    def __init__(
        self, methods: List[Method], report_title: str = "Geotechnical Data Report"
    ):
        self.methods = methods
        self.report_title = report_title
        self.doc = Document()

    def create_report(
        self, description: str, map_image_path: str = r"C:\file\map_screenshot.png"
    ):
        # Add report title
        self.doc.add_heading(self.report_title, 0)

        # Add description text
        self.doc.add_paragraph(description)

        # Add table for methods data (x, y, z, and method name)
        self._add_methods_table()

        # Add the map image to the document
        self.doc.add_picture(map_image_path, width=Inches(6))

        # Save the document
        self.doc.save(f"{self.report_title}.docx")

    def _add_methods_table(self):
        # Create a table with headers
        table = self.doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "X"
        hdr_cells[1].text = "Y"
        hdr_cells[2].text = "Z"
        hdr_cells[3].text = "Method Name"

        # Populate the table with data from Method objects
        for method in self.methods:
            row_cells = table.add_row().cells
            row_cells[0].text = str(method.x)
            row_cells[1].text = str(method.y)
            row_cells[2].text = str(method.z)
            row_cells[3].text = method.name

    def add_existing_map(self, m: folium.Map):
        # Add an existing map to the Word document
        m.save(r"C:\file\methods_map.html")

    def create_folium_map(self, map_path: str = "methods_map.html"):
        # Create a Folium map centered on the average location of all methods
        avg_lat = sum([method.lat for method in self.methods]) / len(self.methods)
        avg_lon = sum([method.lon for method in self.methods]) / len(self.methods)
        folium_map = folium.Map(location=[avg_lat, avg_lon], zoom_start=12)

        # Add markers for each method
        for method in self.methods:
            folium.Marker(
                location=[method.lat, method.lon],
                popup=f"Method: {method.name}, Z: {method.z}",
            ).add_to(folium_map)

        # Save the Folium map as an HTML file
        folium_map.save(map_path)
        return folium_map

    def capture_map_screenshot(
        self,
        map_path: str = r"C:\file\methods_map.html",  # Use a raw string or double backslashes
        screenshot_path: str = r"C:\file\map_screenshot.png",
    ):
        # Use Selenium to take a screenshot of the Folium map
        options = webdriver.ChromeOptions()
        options.add_argument("headless")  # Run Chrome in headless mode
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()), options=options
        )

        # Use 'file://' prefix for accessing local files
        driver.get(f"file://{map_path}")

        # Take a screenshot
        driver.save_screenshot(screenshot_path)

        driver.quit()
